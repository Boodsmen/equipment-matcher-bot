"""Extract text content from DOCX files (paragraphs + tables)."""

from docx import Document
from utils.logger import logger


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text from a DOCX file including paragraphs and tables.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Combined text from all paragraphs and tables.

    Raises:
        ValueError: If file cannot be parsed.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error(f"Failed to open DOCX file {file_path}: {e}")
        raise ValueError(f"Не удалось открыть DOCX файл: {e}") from e

    parts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract tables
    for table in doc.tables:
        table_lines: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells)
            if row_text.replace("|", "").strip():
                table_lines.append(row_text)
        if table_lines:
            parts.append("\n".join(table_lines))

    full_text = "\n\n".join(parts)
    logger.info(f"Extracted {len(full_text)} chars from DOCX ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")
    return full_text
