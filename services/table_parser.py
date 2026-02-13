"""
Direct table-based requirements parser for structured tender documents.

This parser handles documents with structured tables (like Table 2 format):
| Наименование товара | № п/п | Наименование характеристики | Значение | Единица |

Each row becomes a separate requirement, avoiding AI aggregation issues.

Improvements over initial version (inspired by old_bot/processor/parsers/docx_parser.py):
- Dynamic column detection (not fixed positions) using regex header patterns
- Multi-row header support (checks first 3 rows)
- Extracts ALL matching characteristics tables (not just the first one)
- Extracts equipment quantity from a separate equipment-list table
"""

import json
import os
import re
from collections import defaultdict
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from utils.logger import logger


# ═══════════════════════════════════════════════════════════════════════════
# Normalization Map Loader
# ═══════════════════════════════════════════════════════════════════════════


def _load_normalization_map() -> Dict[str, str]:
    """
    Load normalization map and create reverse lookup: variant -> canonical_key.

    Returns:
        Dict mapping lowercase characteristic variants to canonical keys.
    """
    norm_path = os.path.join(
        os.path.dirname(os.path.dirname(__file__)), "data", "normalization_map.json"
    )

    try:
        with open(norm_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        reverse_map = {}
        for canonical_key, variants in data.get("canonical_keys", {}).items():
            for variant in variants:
                clean_variant = variant.lower().strip()
                reverse_map[clean_variant] = canonical_key

        logger.info(f"Loaded normalization map: {len(reverse_map)} variants -> {len(data.get('canonical_keys', {}))} canonical keys")
        return reverse_map

    except Exception as e:
        logger.error(f"Failed to load normalization_map.json: {e}")
        return {}


_NORMALIZATION_MAP = _load_normalization_map()


def normalize_characteristic_name(name: str) -> Optional[str]:
    """
    Normalize a characteristic name to its canonical key.

    Args:
        name: Raw characteristic name from table.

    Returns:
        Canonical key or None if not found.
    """
    clean_name = name.lower().strip()
    return _NORMALIZATION_MAP.get(clean_name)


# ═══════════════════════════════════════════════════════════════════════════
# Value Parsing
# ═══════════════════════════════════════════════════════════════════════════


def parse_value(value_str: str, unit: str = "") -> Any:
    """
    Parse a characteristic value from table cell.

    Handles:
    - Comparison operators: ≥, ≤, >, <, = — СОХРАНЯЮТСЯ в значении
    - Boolean values: Да/Нет, Yes/No
    - Numeric values: 6000, 24, 1.5
    - Text values: "Управляемый", "AC", etc.

    Args:
        value_str: Value from "Значение характеристики" column.
        unit: Unit from "Единица измерения" column (optional).

    Returns:
        Parsed value (int, float, bool, or str).
        Для числовых значений с оператором возвращает строку "<=100", ">=24".
    """
    if not value_str:
        return None

    value_str = value_str.strip()

    # Boolean values
    if value_str.lower() in ["да", "yes", "истина", "true"]:
        return True
    if value_str.lower() in ["нет", "no", "ложь", "false"]:
        return False

    # Определяем наличие оператора сравнения
    operator_match = re.match(r'^([≥≤><≠]=?|>=|<=|!=)\s*', value_str)
    operator = None
    if operator_match:
        raw_op = operator_match.group(1)
        # Нормализуем Unicode операторы в ASCII
        op_map = {'≥': '>=', '≤': '<=', '≠': '!=', '>': '>', '<': '<', '=': '='}
        operator = op_map.get(raw_op, raw_op)

    # Текстовые префиксы → оператор
    if not operator:
        if re.match(r'не\s+менее\b', value_str, re.IGNORECASE):
            operator = ">="
        elif re.match(r'не\s+более\b', value_str, re.IGNORECASE):
            operator = "<="
        elif re.match(r'до\s+', value_str, re.IGNORECASE):
            operator = "<="

    # Извлекаем числа из строки (без оператора)
    value_for_numbers = re.sub(r'^[≥≤><≠=]+\s*', '', value_str)
    value_for_numbers = re.sub(r'^(?:не\s+менее|не\s+более|до)\s+', '', value_for_numbers, flags=re.IGNORECASE)

    numbers = re.findall(r'[\d,]+\.?\d*', value_for_numbers)
    if numbers:
        # Take the last (usually stricter) number
        num_str = numbers[-1].replace(',', '')
        try:
            if '.' not in num_str:
                num_val = int(num_str)
            else:
                num_val = float(num_str)

            # Если есть оператор — возвращаем строку с оператором
            if operator:
                return f"{operator}{num_val}"

            # Чистое число без оператора
            return num_val
        except ValueError:
            pass

    # Return as string if not a number or bool
    return value_str


# ═══════════════════════════════════════════════════════════════════════════
# Column Detection (from old_bot: dynamic, not fixed positions)
# ═══════════════════════════════════════════════════════════════════════════

# Header patterns for characteristics table columns
_ITEM_NAME_PATTERNS = [
    r'наименование\s*(товара|оборудования|изделия|позиции)',
    r'тип\s*(оборудования|устройства)',
    r'раздел',
]
_ITEM_NUMBER_PATTERNS = [
    r'№\s*п/?п',
    r'^п/?п$',
    r'номер',
    r'^\d+$',  # purely numeric header (rare)
]
_CHAR_NAME_PATTERNS = [
    r'наименование\s*характеристик',
    r'характеристик',
    r'параметр',
    r'требование',
    r'показатель',
]
_VALUE_PATTERNS = [
    r'значение\s*(характеристики|параметра)?',
    r'требуемое\s*значение',
    r'^значение$',
    r'величина',
    r'^value$',
]
_UNIT_PATTERNS = [
    r'единица\s*(измерения)?',
    r'ед\.?\s*изм\.?',
    r'размерность',
]


def _match_any_pattern(text: str, patterns: List[str]) -> bool:
    """Return True if text matches any of the given regex patterns."""
    t = text.lower().strip()
    return any(re.search(p, t) for p in patterns)


def _detect_characteristics_columns(table) -> Optional[Dict[str, Any]]:
    """
    Detect column indices in a characteristics table by scanning header rows.

    Checks up to the first 3 rows for headers (supports multi-row headers).

    Returns:
        Dict with keys: 'item_name', 'item_number', 'char_name', 'value', 'unit',
        'header_rows' (number of header rows to skip).
        Returns None if required columns (char_name + value) not found.
    """
    col_map: Dict[str, Optional[int]] = {
        "item_name": None,
        "item_number": None,
        "char_name": None,
        "value": None,
        "unit": None,
    }
    header_rows = 1

    for row_idx in range(min(3, len(table.rows))):
        cells = [cell.text.strip() for cell in table.rows[row_idx].cells]

        for col_idx, cell_text in enumerate(cells):
            if not cell_text:
                continue

            if col_map["item_name"] is None and _match_any_pattern(cell_text, _ITEM_NAME_PATTERNS):
                col_map["item_name"] = col_idx
                header_rows = max(header_rows, row_idx + 1)

            if col_map["item_number"] is None and _match_any_pattern(cell_text, _ITEM_NUMBER_PATTERNS):
                col_map["item_number"] = col_idx
                header_rows = max(header_rows, row_idx + 1)

            if col_map["char_name"] is None and _match_any_pattern(cell_text, _CHAR_NAME_PATTERNS):
                col_map["char_name"] = col_idx
                header_rows = max(header_rows, row_idx + 1)

            if col_map["value"] is None and _match_any_pattern(cell_text, _VALUE_PATTERNS):
                col_map["value"] = col_idx
                header_rows = max(header_rows, row_idx + 1)

            if col_map["unit"] is None and _match_any_pattern(cell_text, _UNIT_PATTERNS):
                col_map["unit"] = col_idx
                header_rows = max(header_rows, row_idx + 1)

    # char_name + value are mandatory
    if col_map["char_name"] is None or col_map["value"] is None:
        # Fallback: if value column not found, try the column right after char_name
        if col_map["char_name"] is not None and col_map["value"] is None:
            fallback_value = col_map["char_name"] + 1
            if fallback_value < len(table.columns):
                col_map["value"] = fallback_value
                logger.debug("Value column not found, using fallback (char_name + 1)")
        else:
            return None

    col_map["header_rows"] = header_rows
    return col_map


def _get_cell(cells: List[str], idx: Optional[int]) -> str:
    """Safely get cell text by index, returns empty string if out of range or None."""
    if idx is None or idx >= len(cells):
        return ""
    return cells[idx]


# ═══════════════════════════════════════════════════════════════════════════
# Characteristics Table Detection
# ═══════════════════════════════════════════════════════════════════════════


def _is_characteristics_table(table) -> bool:
    """
    Check if a table has the expected characteristics format.

    Looks for columns matching characteristic-name and value patterns
    in the first 3 header rows.

    Args:
        table: python-docx Table object.

    Returns:
        True if table matches characteristics format.
    """
    if not table.rows or len(table.rows) < 2:
        return False

    return _detect_characteristics_columns(table) is not None


# ═══════════════════════════════════════════════════════════════════════════
# Equipment List Table (from old_bot: _extract_equipment_list)
# ═══════════════════════════════════════════════════════════════════════════


def _extract_equipment_list(table) -> Dict[str, int]:
    """
    Extract equipment quantities from an equipment-list table.

    Looks for tables with columns: name + quantity.
    Returns a dict: normalized_item_name -> quantity.

    Adapted from old_bot/processor/parsers/docx_parser.py::_extract_equipment_list.

    Args:
        table: python-docx Table object.

    Returns:
        Dict mapping lowercased item name to integer quantity (default 1).
    """
    if len(table.rows) < 2:
        return {}

    # Check first row for "name" and "quantity" columns
    first_row = [cell.text.strip().lower() for cell in table.rows[0].cells]

    has_name = any(
        "наименование" in cell or "оборудование" in cell or "товар" in cell
        for cell in first_row
    )
    has_qty = any(
        "количество" in cell or "кол-во" in cell or "шт" in cell or "qty" in cell
        for cell in first_row
    )

    if not (has_name and has_qty):
        return {}

    # Find column indices
    name_col: Optional[int] = None
    qty_col: Optional[int] = None
    for idx, cell in enumerate(first_row):
        if name_col is None and ("наименование" in cell or "оборудование" in cell or "товар" in cell):
            name_col = idx
        if qty_col is None and ("количество" in cell or "кол-во" in cell or "шт" in cell or "qty" in cell):
            qty_col = idx

    if name_col is None:
        return {}

    result: Dict[str, int] = {}
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if not any(cells):
            continue

        name = _get_cell(cells, name_col)
        if not name:
            continue

        qty = 1
        if qty_col is not None:
            qty_str = _get_cell(cells, qty_col)
            digits = re.findall(r'\d+', qty_str)
            if digits:
                qty = int(digits[0])

        result[name.lower()] = qty

    logger.debug(f"Equipment list table: found {len(result)} items with quantities")
    return result


def _match_quantity(item_name: str, equipment_list: Dict[str, int]) -> Optional[int]:
    """
    Match an item name against the equipment list to find its quantity.

    Uses substring matching (item_name ⊂ list_name or vice versa).

    Args:
        item_name: Item name from the characteristics table.
        equipment_list: {lowercased_name: quantity} from _extract_equipment_list.

    Returns:
        Quantity if found, None otherwise.
    """
    name_lower = item_name.lower().strip()
    for list_name, qty in equipment_list.items():
        if name_lower in list_name or list_name in name_lower:
            return qty
    return None


# ═══════════════════════════════════════════════════════════════════════════
# Table Parsing
# ═══════════════════════════════════════════════════════════════════════════


def _parse_table_rows(table, col_map: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Parse all data rows from a characteristics table using detected column map.

    Args:
        table: python-docx Table object.
        col_map: Column map from _detect_characteristics_columns.

    Returns:
        List of parsed row dicts.
    """
    header_rows: int = col_map.get("header_rows", 1)
    parsed_rows = []

    # Track last seen item_name and item_number for merged cells
    last_item_name = ""
    last_item_number = ""

    for row in table.rows[header_rows:]:
        cells = [cell.text.strip() for cell in row.cells]

        if not any(cells):
            continue

        # Get cell values using detected column indices
        item_name = _get_cell(cells, col_map["item_name"]) or last_item_name
        item_number = _get_cell(cells, col_map["item_number"]) or last_item_number
        char_name = _get_cell(cells, col_map["char_name"])
        value = _get_cell(cells, col_map["value"])
        unit = _get_cell(cells, col_map["unit"])

        # Update last seen values for merged-cell propagation
        if _get_cell(cells, col_map["item_name"]):
            last_item_name = item_name
        if _get_cell(cells, col_map["item_number"]):
            last_item_number = item_number

        # Skip rows without characteristic name
        if not char_name:
            continue

        # Skip rows that look like sub-headers
        if _match_any_pattern(char_name, _CHAR_NAME_PATTERNS):
            continue

        # Normalize characteristic name
        canonical_key = normalize_characteristic_name(char_name)
        if not canonical_key:
            logger.debug(f"Could not normalize characteristic: '{char_name}'")
            # Fallback: snake_case from raw name
            canonical_key = re.sub(r'\W+', '_', char_name.lower().strip()).strip('_')

        # Parse value
        parsed_value = parse_value(value, unit)

        parsed_rows.append({
            "item_name": item_name,
            "item_number": item_number,
            "characteristic_name": char_name,
            "canonical_key": canonical_key,
            "value": value,
            "unit": unit,
            "parsed_value": parsed_value,
        })

    return parsed_rows


def _group_requirements_by_item(rows: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
    """
    Group parsed rows by equipment item.

    Uses item_number prefix (e.g., "1.1", "1.2" -> item 1, "2.1", "2.2" -> item 2).
    Falls back to item_name if item_number is missing.

    Args:
        rows: List of parsed row dicts.

    Returns:
        Dict mapping group_key -> list of requirements.
    """
    groups: Dict[str, List[Dict[str, Any]]] = defaultdict(list)

    for row in rows:
        item_num = row["item_number"]

        # Extract numeric prefix (e.g., "1.1" -> "1", "2.3" -> "2")
        match = re.match(r'^(\d+)', item_num)
        if match:
            prefix = match.group(1)
        elif row["item_name"]:
            # No item_number — group by item_name
            prefix = row["item_name"]
        else:
            prefix = "default"

        groups[prefix].append(row)

    return groups


def _build_item_dict(
    item_prefix: str,
    requirements: List[Dict[str, Any]],
    equipment_list: Dict[str, int],
) -> Optional[Dict[str, Any]]:
    """
    Build a single item dict from grouped requirements.

    Args:
        item_prefix: Item number prefix (e.g., "1", "2") or item_name.
        requirements: List of requirement dicts for this item.
        equipment_list: Equipment list {lowercased_name: quantity} for quantity lookup.

    Returns:
        Dict compatible with OpenAI format:
        {
            "item_name": "Коммутатор (позиция 1)",
            "quantity": 5,
            "model_name": null,
            "category": "Коммутаторы",
            "required_specs": {...}
        }
    """
    if not requirements:
        return None

    # Use most common non-empty item_name from requirements
    item_names = [r["item_name"] for r in requirements if r["item_name"]]
    item_name = item_names[0] if item_names else f"Позиция {item_prefix}"

    # Try to infer category from item_name
    category = None
    item_lower = item_name.lower()
    if "коммутатор" in item_lower or "switch" in item_lower:
        category = "Коммутаторы"
    elif "маршрутизатор" in item_lower or "router" in item_lower:
        category = "Маршрутизаторы"

    # Build required_specs dict
    required_specs: Dict[str, Any] = {}
    for req in requirements:
        canonical_key = req["canonical_key"]
        parsed_value = req["parsed_value"]
        if parsed_value is not None:
            required_specs[canonical_key] = parsed_value

    # Try to get quantity from equipment list table
    quantity = _match_quantity(item_name, equipment_list)

    return {
        "item_name": f"{item_name} (позиция {item_prefix})" if item_prefix.isdigit() else item_name,
        "quantity": quantity,
        "model_name": None,
        "category": category,
        "required_specs": required_specs,
    }


# ═══════════════════════════════════════════════════════════════════════════
# Main Parser
# ═══════════════════════════════════════════════════════════════════════════


def parse_requirements_from_tables(file_path: str) -> Optional[Dict[str, Any]]:
    """
    Parse requirements from structured tables in a DOCX file.

    This is the main entry point. Returns None if no suitable tables found.

    Improvements over v1:
    - Dynamic column detection (not fixed positions)
    - Scans ALL tables in document (not just first matching one)
    - Extracts quantities from equipment-list table
    - Propagates merged cell values (item_name / item_number)

    Args:
        file_path: Path to DOCX file.

    Returns:
        Dict with 'items' list (compatible with OpenAI format), or None if parsing failed.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error(f"Failed to open DOCX: {e}")
        return None

    logger.info(f"Analyzing {len(doc.tables)} tables in document")

    # Pass 1: scan all tables — classify each as characteristics table or equipment list
    characteristics_tables: List[Tuple[int, Any, Dict]] = []  # (idx, table, col_map)
    equipment_list: Dict[str, int] = {}

    for idx, table in enumerate(doc.tables):
        col_map = _detect_characteristics_columns(table)
        if col_map is not None:
            logger.info(f"Found characteristics table at index {idx} ({len(table.rows)} rows), columns={col_map}")
            characteristics_tables.append((idx, table, col_map))
        else:
            # Try as equipment list table
            eq_list = _extract_equipment_list(table)
            if eq_list:
                logger.info(f"Found equipment list table at index {idx} ({len(eq_list)} items)")
                equipment_list.update(eq_list)

    if not characteristics_tables:
        logger.info("No structured characteristics table found")
        return None

    if equipment_list:
        logger.info(f"Equipment list: {dict(list(equipment_list.items())[:5])}{'...' if len(equipment_list) > 5 else ''}")

    # Pass 2: parse all characteristics tables and merge rows
    all_parsed_rows: List[Dict[str, Any]] = []
    for idx, table, col_map in characteristics_tables:
        rows = _parse_table_rows(table, col_map)
        logger.info(f"Table {idx}: parsed {len(rows)} requirement rows")
        all_parsed_rows.extend(rows)

    logger.info(f"Total parsed rows: {len(all_parsed_rows)}")

    if not all_parsed_rows:
        return None

    # Group by item
    groups = _group_requirements_by_item(all_parsed_rows)
    logger.info(f"Grouped into {len(groups)} equipment items")

    # Build items list
    items = []
    for prefix in sorted(groups.keys(), key=lambda x: int(x) if x.isdigit() else 0):
        item_dict = _build_item_dict(prefix, groups[prefix], equipment_list)
        if item_dict:
            items.append(item_dict)

    result = {"items": items}

    logger.info(
        f"Table parser extracted {len(items)} items, "
        f"total specs: {sum(len(item['required_specs']) for item in items)}"
    )

    return result
